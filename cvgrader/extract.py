"""Stage 1: text extraction and layout-hazard detection.

This simulates the parsing engines (Textkernel, Sovren/Aqore, Daxtra, HireAbility)
that sit inside nearly every commercial ATS. Most "the ATS rejected me" stories
are really "the parser scrambled my resume" stories, so hazards found here are
weighted heavily in the final score.
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

_EMAIL_RE = re.compile(r"[\w.+-]+@[\w-]+\.[\w.-]+")
_PHONEISH_RE = re.compile(r"\+?\d[\d\s().\-]{7,16}\d")


@dataclass
class Hazard:
    code: str
    message: str
    fix: str
    severe: bool = False


@dataclass
class Extraction:
    path: str
    filetype: str
    text: str = ""
    page_count: int = 1
    hazards: list[Hazard] = field(default_factory=list)
    error: str = ""

    @property
    def words(self) -> int:
        return len(self.text.split())

    @property
    def lines(self) -> list[str]:
        return self.text.splitlines()


def extract(path) -> Extraction:
    p = Path(path)
    if not p.exists():
        return Extraction(str(p), "missing", error=f"File not found: {p}")
    ext = p.suffix.lower()
    try:
        if ext == ".pdf":
            return _from_pdf(p)
        if ext == ".docx":
            return _from_docx(p)
        if ext in (".txt", ".md"):
            text = p.read_text(encoding="utf-8", errors="replace")
            ex = Extraction(str(p), ext.lstrip("."), text)
            ex.page_count = max(1, round(ex.words / 550))
            return ex
    except Exception as e:
        return Extraction(str(p), ext.lstrip("."), error=f"Parser error ({type(e).__name__}): {e}")
    if ext == ".doc":
        return Extraction(str(p), "doc", error=(
            "Legacy .doc format - several ATS parsers mangle it. Save as .docx or PDF."))
    return Extraction(str(p), ext.lstrip(".") or "unknown",
                      error=f"Unsupported file type '{ext}'. Use PDF, DOCX or TXT.")


def _from_pdf(p: Path) -> Extraction:
    import pdfplumber

    ex = Extraction(str(p), "pdf")
    texts: list[str] = []
    image_count = 0
    table_pages = 0
    column_pages = 0
    first_lines: list[str] = []
    last_lines: list[str] = []

    with pdfplumber.open(str(p)) as pdf:
        ex.page_count = len(pdf.pages)
        for page in pdf.pages:
            text = page.extract_text() or ""
            texts.append(text)
            image_count += len(page.images)
            try:
                if page.find_tables():
                    table_pages += 1
            except Exception:
                pass
            if _looks_multicolumn(page):
                column_pages += 1
            page_lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
            if page_lines:
                first_lines.append(page_lines[0])
                last_lines.append(page_lines[-1])

    ex.text = "\n".join(texts)

    if ex.words < 40:
        if image_count:
            ex.hazards.append(Hazard(
                "scanned_pdf",
                "PDF has images but almost no extractable text - it is likely a scanned "
                "image or an export that flattened text into graphics.",
                "Re-export the resume so the text layer is selectable (try Ctrl+A in a "
                "PDF viewer: if you cannot select the text, neither can an ATS).",
                severe=True))
        else:
            ex.hazards.append(Hazard(
                "no_text", "Almost no text could be extracted from this PDF.",
                "Re-export from your editor as a text-based PDF.", severe=True))
    elif image_count:
        ex.hazards.append(Hazard(
            "images",
            f"{image_count} embedded image(s) found. Photos, logos and icon fonts are "
            "invisible to parsers, and headshots invite bias screening at some companies.",
            "Remove photos/graphics; replace skill icons or rating bars with plain text."))

    if table_pages:
        ex.hazards.append(Hazard(
            "tables",
            f"Table layout detected on {table_pages} page(s). Many parsers read tables "
            "cell-by-cell in the wrong order, scrambling your work history.",
            "Rebuild tabular sections as plain left-aligned lines."))
    if column_pages:
        ex.hazards.append(Hazard(
            "multi_column",
            f"Multi-column layout detected on {column_pages} page(s). Parsers read "
            "straight across the page, interleaving the columns into nonsense.",
            "Switch to a single-column layout."))

    if ex.page_count >= 2:
        repeated = {ln for ln in first_lines if len(ln) > 3 and first_lines.count(ln) >= 2}
        repeated |= {ln for ln in last_lines if len(ln) > 3 and last_lines.count(ln) >= 2}
        if repeated:
            sample = next(iter(repeated))
            has_contact = bool(_EMAIL_RE.search(" ".join(repeated))
                               or _PHONEISH_RE.search(" ".join(repeated)))
            ex.hazards.append(Hazard(
                "header_footer",
                f"Repeated running header/footer detected ('{sample[:60]}'). "
                + ("It contains contact info, which some parsers drop entirely."
                   if has_contact else "Parsers may merge it into your content."),
                "Put contact details in the document body, not the header/footer.",
                severe=has_contact))
    return ex


def _looks_multicolumn(page) -> bool:
    """Line-based heuristic: in a 2-column layout most text lines live entirely
    in one half of the page; in single-column text most lines span the center."""
    try:
        words = page.extract_words()
    except Exception:
        return False
    if len(words) < 60:
        return False
    cx = page.width / 2
    rows: dict[int, list] = {}
    for w in words:
        rows.setdefault(int(w["top"] // 4), []).append(w)
    spanning = confined = 0
    for ws in rows.values():
        if len(ws) < 2:
            continue
        has_left = any(w["x0"] < cx * 0.85 for w in ws)
        has_right = any(w["x1"] > cx * 1.15 for w in ws)
        if has_left and has_right:
            spanning += 1
        else:
            confined += 1
    return confined >= 10 and confined > 3 * spanning


def _from_docx(p: Path) -> Extraction:
    import docx
    from docx.oxml.ns import qn

    doc = docx.Document(str(p))
    ex = Extraction(str(p), "docx")
    parts = [para.text for para in doc.paragraphs]

    table_count = len(doc.tables)
    for tb in doc.tables:
        for row in tb.rows:
            seen = []
            for cell in row.cells:
                t = cell.text.strip()
                if t and (not seen or seen[-1] != t):  # merged cells repeat text
                    seen.append(t)
            if seen:
                parts.append("  ".join(seen))

    # Text boxes: invisible to python-docx paragraphs AND to many ATS parsers.
    textbox_nodes = doc.element.body.findall(".//" + qn("w:txbxContent"))
    textbox_words = 0
    for node in textbox_nodes:
        tb_text = " ".join(t.text or "" for t in node.findall(".//" + qn("w:t")))
        textbox_words += len(tb_text.split())
        if tb_text.strip():
            parts.append(tb_text.strip())

    header_footer_text = ""
    for sec in doc.sections:
        for container in (sec.header, sec.footer):
            for para in container.paragraphs:
                header_footer_text += " " + para.text
    if header_footer_text.strip():
        parts.append(header_footer_text.strip())

    ex.text = "\n".join(parts)
    ex.page_count = max(1, round(ex.words / 550))

    if table_count:
        ex.hazards.append(Hazard(
            "tables",
            f"{table_count} Word table(s) found. Parsers often read table cells in the "
            "wrong order or drop them.",
            "Rebuild tabular sections as plain left-aligned lines (use tab stops or "
            "spacing, not table cells)."))
    if textbox_nodes:
        ex.hazards.append(Hazard(
            "text_boxes",
            f"{len(textbox_nodes)} text box(es) holding ~{textbox_words} words. Many "
            "parsers skip text boxes completely - that content may not exist to an ATS.",
            "Move all content out of text boxes into normal paragraphs.",
            severe=textbox_words > 30))
    if len(doc.inline_shapes):
        ex.hazards.append(Hazard(
            "images",
            f"{len(doc.inline_shapes)} embedded image(s). Invisible to parsers.",
            "Remove photos/graphics; keep everything as text."))
    if header_footer_text.strip():
        has_contact = bool(_EMAIL_RE.search(header_footer_text)
                           or _PHONEISH_RE.search(header_footer_text))
        ex.hazards.append(Hazard(
            "header_footer",
            "Content found in the Word header/footer"
            + (" including contact info - several parsers drop headers, which means "
               "no email/phone gets imported." if has_contact else "."),
            "Put your name and contact details in the document body.",
            severe=has_contact))
    return ex
